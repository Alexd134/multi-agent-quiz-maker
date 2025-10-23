"""DOCX document generator for quiz export."""

import os
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from src.models.quiz import Quiz, QuizRound


def ensure_output_directory(output_dir: str = "output") -> Path:
    """
    Ensure the output directory exists.

    Args:
        output_dir: Directory path to create

    Returns:
        Path object for the output directory
    """
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    return output_path


def generate_timestamped_filename(base_name: str, extension: str = "docx") -> str:
    """
    Generate a filename with timestamp.

    Args:
        base_name: Base name for the file
        extension: File extension (without dot)

    Returns:
        Filename with timestamp
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    # Clean the base name to remove any path components
    base_name = Path(base_name).name
    return f"{base_name}_{timestamp}.{extension}"


def export_to_docx(
    quiz: Quiz,
    output_path: str,
    include_answers: bool = False,
    use_output_dir: bool = True,
    output_dir: str = "output"
) -> str:
    """
    Export quiz to a formatted DOCX file.

    Args:
        quiz: Quiz object to export
        output_path: Path where the DOCX file should be saved (can be relative or absolute)
        include_answers: If True, includes correct answers and explanations
        use_output_dir: If True, saves to output directory with timestamp (default: True)
        output_dir: Directory to save files in (default: "output")

    Returns:
        Path to the created DOCX file
    """
    # If use_output_dir is True, modify the output path
    if use_output_dir:
        output_dir_path = ensure_output_directory(output_dir)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = Path(output_path).stem  # Get name without extension
        filename = f"{base_name}_{timestamp}.docx"
        output_path = str(output_dir_path / filename)

    # Create document
    doc = Document()

    # Set up styles
    setup_document_styles(doc)

    # Add title
    title = doc.add_heading(quiz.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add description if present
    if quiz.description:
        desc_para = doc.add_paragraph(quiz.description)
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_para.runs[0].italic = True

    # Add quiz info
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.add_run(f"Total Rounds: {quiz.total_rounds}").bold = True
    info_para.add_run(f"  |  ")
    info_para.add_run(f"Total Questions: {quiz.total_questions}").bold = True
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add generation date
    date_para = doc.add_paragraph(
        f"Generated: {quiz.metadata.created_at.strftime('%Y-%m-%d %H:%M')}"
    )
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(9)
    date_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Add page break before questions
    doc.add_page_break()

    # Add each round
    for round_data in quiz.rounds:
        add_round_to_document(doc, round_data, include_answers)
        doc.add_page_break()

    # Add answer key at the end if requested
    if include_answers:
        add_answer_key(doc, quiz)

    # Save document
    doc.save(output_path)

    return output_path


def setup_document_styles(doc: Document) -> None:
    """
    Set up document-wide styles.

    Args:
        doc: Document to configure
    """
    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_round_to_document(
    doc: Document, round_data: QuizRound, include_answers: bool = False
) -> None:
    """
    Add a round to the document.

    Args:
        doc: Document to add to
        round_data: QuizRound object
        include_answers: If True, includes answers and explanations
    """
    # Round header
    round_heading = doc.add_heading(
        f"Round {round_data.round_number}: {round_data.round_name}", level=1
    )
    round_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    # Round info
    info_para = doc.add_paragraph()
    info_para.add_run(f"Topic: {round_data.topic}").italic = True
    info_para.add_run(f"  |  ")
    info_para.add_run(f"Questions: {round_data.question_count}").italic = True

    doc.add_paragraph()

    # Add questions
    for i, question in enumerate(round_data.questions, 1):
        # Question number and text
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{i}. ")
        q_run.bold = True
        q_run.font.size = Pt(12)
        q_para.add_run(question.question_text)

        # Add difficulty indicator
        difficulty_para = doc.add_paragraph()
        difficulty_run = difficulty_para.add_run(
            f"  Difficulty: {question.difficulty.value.capitalize()}"
        )
        difficulty_run.font.size = Pt(9)
        difficulty_run.italic = True

        # Set color based on difficulty
        if question.difficulty.value == "easy":
            difficulty_run.font.color.rgb = RGBColor(0, 128, 0)
        elif question.difficulty.value == "medium":
            difficulty_run.font.color.rgb = RGBColor(255, 140, 0)
        else:  # hard
            difficulty_run.font.color.rgb = RGBColor(255, 0, 0)

        # Options
        options = ["A", "B", "C", "D"]
        for option in options:
            opt_para = doc.add_paragraph(f"   {option}. {question.options[option]}")
            opt_para.paragraph_format.left_indent = Inches(0.5)

            # Highlight correct answer if include_answers is True
            if include_answers and option == question.correct_answer:
                opt_para.runs[0].bold = True
                opt_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)
                opt_para.add_run(" ✓").font.color.rgb = RGBColor(0, 128, 0)

        # Add explanation if include_answers and explanation exists
        if include_answers and question.explanation:
            exp_para = doc.add_paragraph()
            exp_para.paragraph_format.left_indent = Inches(0.5)
            exp_run = exp_para.add_run(f"Explanation: {question.explanation}")
            exp_run.italic = True
            exp_run.font.size = Pt(10)
            exp_run.font.color.rgb = RGBColor(64, 64, 64)

        # Add spacing between questions
        doc.add_paragraph()


def add_answer_key(doc: Document, quiz: Quiz) -> None:
    """
    Add an answer key section at the end of the document.

    Args:
        doc: Document to add to
        quiz: Quiz object
    """
    # Add page break before answer key
    doc.add_page_break()

    # Answer key header
    header = doc.add_heading("Answer Key", level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    # Add answers by round
    for round_data in quiz.rounds:
        # Round header
        round_heading = doc.add_heading(
            f"Round {round_data.round_number}: {round_data.round_name}", level=2
        )
        round_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

        # Create table for answers
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Q#"
        header_cells[1].text = "Answer"
        header_cells[2].text = "Explanation"

        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add answers
        for i, question in enumerate(round_data.questions, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = f"{question.correct_answer} - {question.options[question.correct_answer]}"
            row_cells[2].text = question.explanation or "N/A"

        doc.add_paragraph()


def generate_answer_key(quiz: Quiz, output_path: str) -> str:
    """
    Generate a separate answer key document.

    Args:
        quiz: Quiz object
        output_path: Path where the answer key should be saved

    Returns:
        Path to the created answer key file
    """
    # Create document
    doc = Document()
    setup_document_styles(doc)

    # Title
    title = doc.add_heading(f"{quiz.title} - Answer Key", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Add answer key
    add_answer_key(doc, quiz)

    # Save
    doc.save(output_path)

    return output_path


def export_quiz_with_separate_answers(
    quiz: Quiz, base_path: str, output_dir: str = "output"
) -> tuple[str, str]:
    """
    Export quiz with questions and answers in separate files.

    Args:
        quiz: Quiz object
        base_path: Base path for output files (without extension)
        output_dir: Directory to save files in (default: "output")

    Returns:
        Tuple of (questions_path, answers_path)
    """
    # Ensure output directory exists
    output_path = ensure_output_directory(output_dir)

    # Generate timestamped filenames
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = Path(base_path).name

    questions_filename = f"{base_name}_questions_{timestamp}.docx"
    answers_filename = f"{base_name}_answers_{timestamp}.docx"

    questions_path = str(output_path / questions_filename)
    answers_path = str(output_path / answers_filename)

    # Export questions without answers (use_output_dir=False since we already handled paths)
    export_to_docx(quiz, questions_path, include_answers=False, use_output_dir=False)

    # Export answer key
    generate_answer_key(quiz, answers_path)

    return questions_path, answers_path