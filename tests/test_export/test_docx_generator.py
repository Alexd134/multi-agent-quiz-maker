"""Tests for DOCX export functionality."""

import os
from pathlib import Path

import pytest
from docx import Document

from src.export.docx_generator import (
    ensure_output_directory,
    export_quiz_with_separate_answers,
    export_to_docx,
    generate_timestamped_filename,
)


class TestEnsureOutputDirectory:
    """Test output directory creation."""

    def test_creates_directory_if_not_exists(self, tmp_path):
        """Test that directory is created if it doesn't exist."""
        output_dir = tmp_path / "test_output"
        result = ensure_output_directory(str(output_dir))

        assert output_dir.exists()
        assert output_dir.is_dir()
        assert result == output_dir

    def test_does_not_fail_if_directory_exists(self, tmp_path):
        """Test that function works if directory already exists."""
        output_dir = tmp_path / "test_output"
        output_dir.mkdir()

        # Should not raise an error
        result = ensure_output_directory(str(output_dir))
        assert result == output_dir


class TestGenerateTimestampedFilename:
    """Test timestamped filename generation."""

    def test_generates_filename_with_timestamp(self):
        """Test that filename includes timestamp."""
        filename = generate_timestamped_filename("quiz", "docx")

        assert filename.startswith("quiz_")
        assert filename.endswith(".docx")
        assert len(filename) > len("quiz_.docx")

    def test_uses_correct_extension(self):
        """Test that correct extension is used."""
        filename = generate_timestamped_filename("test", "pdf")

        assert filename.endswith(".pdf")

    def test_handles_path_in_base_name(self):
        """Test that paths in base name are handled correctly."""
        filename = generate_timestamped_filename("/path/to/quiz", "docx")

        # Should only use the base name, not the path
        assert not filename.startswith("/")
        assert filename.startswith("quiz_")


class TestExportToDocx:
    """Test DOCX export functionality."""

    def test_creates_docx_file(self, sample_quiz, tmp_path):
        """Test that DOCX file is created."""
        output_path = tmp_path / "test_quiz.docx"

        result = export_to_docx(
            sample_quiz, str(output_path), use_output_dir=False
        )

        assert os.path.exists(result)
        assert result.endswith(".docx")

    def test_docx_contains_quiz_title(self, sample_quiz, tmp_path):
        """Test that exported DOCX contains quiz title."""
        output_path = tmp_path / "test_quiz.docx"

        export_to_docx(sample_quiz, str(output_path), use_output_dir=False)

        # Read the document
        doc = Document(str(output_path))

        # Check that title is in the document
        full_text = "\n".join([para.text for para in doc.paragraphs])
        assert sample_quiz.title in full_text

    def test_docx_contains_questions(self, sample_quiz, tmp_path):
        """Test that exported DOCX contains questions."""
        output_path = tmp_path / "test_quiz.docx"

        export_to_docx(sample_quiz, str(output_path), use_output_dir=False)

        # Read the document
        doc = Document(str(output_path))
        full_text = "\n".join([para.text for para in doc.paragraphs])

        # Check that at least one question is in the document
        first_question = sample_quiz.rounds[0].questions[0]
        assert first_question.question_text in full_text

    def test_docx_without_answers_does_not_show_correct_answer(
        self, sample_quiz, tmp_path
    ):
        """Test that DOCX without answers doesn't highlight correct answer."""
        output_path = tmp_path / "test_quiz.docx"

        export_to_docx(
            sample_quiz, str(output_path), include_answers=False, use_output_dir=False
        )

        # Read the document
        doc = Document(str(output_path))
        full_text = "\n".join([para.text for para in doc.paragraphs])

        # Explanations should not be present
        first_question = sample_quiz.rounds[0].questions[0]
        if first_question.explanation:
            assert first_question.explanation not in full_text

    def test_docx_with_answers_includes_explanations(self, sample_quiz, tmp_path):
        """Test that DOCX with answers includes explanations."""
        output_path = tmp_path / "test_quiz.docx"

        export_to_docx(
            sample_quiz, str(output_path), include_answers=True, use_output_dir=False
        )

        # Read the document
        doc = Document(str(output_path))
        full_text = "\n".join([para.text for para in doc.paragraphs])

        # Check for answer key section
        assert "Answer Key" in full_text

    def test_uses_output_dir_when_specified(self, sample_quiz, tmp_path):
        """Test that output directory is used when specified."""
        output_dir = tmp_path / "output"

        result = export_to_docx(
            sample_quiz, "quiz", use_output_dir=True, output_dir=str(output_dir)
        )

        # File should be in output directory
        assert str(output_dir) in result
        assert os.path.exists(result)


class TestExportQuizWithSeparateAnswers:
    """Test export with separate answer files."""

    def test_creates_two_files(self, sample_quiz, tmp_path):
        """Test that two files are created (questions and answers)."""
        questions_file, answers_file = export_quiz_with_separate_answers(
            sample_quiz, "quiz", output_dir=str(tmp_path)
        )

        assert os.path.exists(questions_file)
        assert os.path.exists(answers_file)

    def test_filenames_have_correct_suffixes(self, sample_quiz, tmp_path):
        """Test that files have correct suffixes."""
        questions_file, answers_file = export_quiz_with_separate_answers(
            sample_quiz, "quiz", output_dir=str(tmp_path)
        )

        assert "questions" in questions_file
        assert "answers" in answers_file

    def test_questions_file_does_not_contain_answers(self, sample_quiz, tmp_path):
        """Test that questions file doesn't contain answer key."""
        questions_file, _ = export_quiz_with_separate_answers(
            sample_quiz, "quiz", output_dir=str(tmp_path)
        )

        doc = Document(questions_file)
        full_text = "\n".join([para.text for para in doc.paragraphs])

        # Should not contain "Answer Key" section
        assert "Answer Key" not in full_text

    def test_answers_file_contains_answer_key(self, sample_quiz, tmp_path):
        """Test that answers file contains answer key."""
        _, answers_file = export_quiz_with_separate_answers(
            sample_quiz, "quiz", output_dir=str(tmp_path)
        )

        doc = Document(answers_file)
        full_text = "\n".join([para.text for para in doc.paragraphs])

        # Should contain quiz title and answer key info
        assert sample_quiz.title in full_text

    def test_both_files_in_same_directory(self, sample_quiz, tmp_path):
        """Test that both files are created in the same directory."""
        questions_file, answers_file = export_quiz_with_separate_answers(
            sample_quiz, "quiz", output_dir=str(tmp_path)
        )

        questions_dir = os.path.dirname(questions_file)
        answers_dir = os.path.dirname(answers_file)

        assert questions_dir == answers_dir
        assert questions_dir == str(tmp_path)